from __future__ import annotations

import argparse
import hashlib
import logging
import re
from dataclasses import dataclass
from logging.handlers import RotatingFileHandler
from pathlib import Path
from typing import Iterable, List, Optional

import chromadb
import pandas as pd
from chromadb.config import Settings
from docx import Document
from dotenv import load_dotenv
from sentence_transformers import SentenceTransformer

try:
    import fitz
except ImportError:
    fitz = None

try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except ImportError:
    RecursiveCharacterTextSplitter = None

BASE_DIR = Path(__file__).resolve().parent
KNOWLEDGE_BASE_DIR = BASE_DIR / "Phase 1 (RAG)"
UPLOADS_DIR = BASE_DIR / "uploads"
CHROMA_DIR = BASE_DIR / "chroma_db"
LOG_DIR = BASE_DIR / "logs"
COLLECTION_NAME = "recycling"
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".csv", ".xlsx"}
EMBEDDING_MODEL_NAME = "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"

load_dotenv(BASE_DIR / ".env")


def setup_logger() -> logging.Logger:
    LOG_DIR.mkdir(exist_ok=True)
    logger = logging.getLogger("notun_alo_rag")
    logger.setLevel(logging.INFO)
    logger.propagate = False
    if not logger.handlers:
        handler = RotatingFileHandler(LOG_DIR / "rag.log", maxBytes=1_000_000, backupCount=5, encoding="utf-8")
        handler.setFormatter(logging.Formatter("%(asctime)s [%(levelname)s] %(message)s"))
        logger.addHandler(handler)
        console = logging.StreamHandler()
        console.setFormatter(logging.Formatter("%(levelname)s: %(message)s"))
        logger.addHandler(console)
    return logger


logger = setup_logger()


@dataclass
class LoadedDocument:
    text: str
    source_path: Path
    page_number: int


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"\s+", " ", text, flags=re.UNICODE)
    return text.strip()


def file_hash(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def folder_category(path: Path) -> str:
    try:
        rel = path.relative_to(KNOWLEDGE_BASE_DIR)
        return rel.parts[0] if len(rel.parts) > 1 else "root"
    except ValueError:
        return "uploads"


def load_pdf(path: Path) -> List[LoadedDocument]:
    if fitz is None:
        raise RuntimeError("PyMuPDF is not installed. Run: pip install -r requirements.txt")
    docs: List[LoadedDocument] = []
    with fitz.open(path) as pdf:
        for page_index, page in enumerate(pdf, start=1):
            text = clean_text(page.get_text("text"))
            if text:
                docs.append(LoadedDocument(text=text, source_path=path, page_number=page_index))
    return docs


def load_docx(path: Path) -> List[LoadedDocument]:
    doc = Document(path)
    parts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    text = clean_text("\n".join(parts))
    return [LoadedDocument(text=text, source_path=path, page_number=1)] if text else []


def load_txt(path: Path) -> List[LoadedDocument]:
    text = clean_text(path.read_text(encoding="utf-8", errors="ignore"))
    return [LoadedDocument(text=text, source_path=path, page_number=1)] if text else []


def load_table(path: Path) -> List[LoadedDocument]:
    if path.suffix.lower() == ".csv":
        try:
            df = pd.read_csv(path, encoding_errors="ignore")
        except Exception:
            logger.warning("Retrying irregular CSV with tolerant parser: %s", path.name)
            df = pd.read_csv(
                path,
                encoding_errors="ignore",
                engine="python",
                sep=None,
                on_bad_lines="skip",
            )
    else:
        df = pd.read_excel(path)
    df = df.fillna("")
    text = clean_text(df.to_csv(index=False))
    return [LoadedDocument(text=text, source_path=path, page_number=1)] if text else []


def load_file(path: Path) -> List[LoadedDocument]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return load_pdf(path)
    if suffix == ".docx":
        return load_docx(path)
    if suffix == ".txt":
        return load_txt(path)
    if suffix in {".csv", ".xlsx"}:
        return load_table(path)
    return []


def iter_supported_files(extra_files: Optional[Iterable[Path]] = None) -> Iterable[Path]:
    for base in [KNOWLEDGE_BASE_DIR, UPLOADS_DIR]:
        if base.exists():
            for path in base.rglob("*"):
                if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
                    yield path
    if extra_files:
        for path in extra_files:
            path = Path(path)
            if path.exists() and path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
                yield path


def get_collection(rebuild: bool = False):
    CHROMA_DIR.mkdir(exist_ok=True)
    client = chromadb.PersistentClient(path=str(CHROMA_DIR), settings=Settings(anonymized_telemetry=False))
    if rebuild:
        try:
            client.delete_collection(COLLECTION_NAME)
            logger.info("Deleted existing Chroma collection '%s'", COLLECTION_NAME)
        except Exception:
            logger.info("No existing Chroma collection to delete")
    return client.get_or_create_collection(COLLECTION_NAME)


def make_splitter():
    if RecursiveCharacterTextSplitter is None:
        class SimpleSplitter:
            def split_text(self, text: str) -> List[str]:
                chunks = []
                for start in range(0, len(text), 800):
                    chunk = text[start:start + 1000].strip()
                    if chunk:
                        chunks.append(chunk)
                return chunks
        return SimpleSplitter()
    return RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)


def is_indexed(collection, digest: str) -> bool:
    try:
        result = collection.get(where={"file_hash": digest}, limit=1)
        return bool(result.get("ids"))
    except Exception as exc:
        logger.warning("Could not check index state for hash %s: %s", digest[:10], exc)
        return False


def ingest_documents(rebuild: bool = False, extra_files: Optional[Iterable[Path]] = None, only_file: Optional[Path] = None) -> dict:
    collection = get_collection(rebuild=rebuild)
    splitter = make_splitter()
    model = SentenceTransformer(EMBEDDING_MODEL_NAME, local_files_only=True)
    files = [Path(only_file)] if only_file else list(dict.fromkeys(iter_supported_files(extra_files)))
    indexed_files = skipped_files = failed_files = total_chunks = 0
    logger.info("Starting ingestion. rebuild=%s files=%s", rebuild, len(files))

    for path in files:
        try:
            if not path.exists() or path.suffix.lower() not in SUPPORTED_EXTENSIONS:
                skipped_files += 1
                logger.info("Skipping unsupported or missing file: %s", path)
                continue
            digest = file_hash(path)
            if not rebuild and is_indexed(collection, digest):
                skipped_files += 1
                logger.info("Skipping already indexed file: %s", path.name)
                continue
            loaded_docs = load_file(path)
            if not loaded_docs:
                skipped_files += 1
                logger.info("Skipping empty file after extraction: %s", path.name)
                continue
            ids, texts, metadatas = [], [], []
            chunk_counter = 0
            for doc in loaded_docs:
                for chunk in splitter.split_text(doc.text):
                    chunk_counter += 1
                    chunk_id = f"{digest[:16]}-{doc.page_number}-{chunk_counter}"
                    ids.append(chunk_id)
                    texts.append(chunk)
                    metadatas.append({
                        "chunk_id": chunk_id,
                        "filename": path.name,
                        "page_number": int(doc.page_number),
                        "source_doc": str(path),
                        "folder_category": folder_category(path),
                        "file_type": path.suffix.lower().lstrip("."),
                        "file_hash": digest,
                    })
            if not texts:
                skipped_files += 1
                logger.info("No chunks produced for file: %s", path.name)
                continue
            logger.info("Generating embeddings for %s chunks from %s", len(texts), path.name)
            embeddings = model.encode(texts, show_progress_bar=False).tolist()
            collection.add(ids=ids, documents=texts, metadatas=metadatas, embeddings=embeddings)
            indexed_files += 1
            total_chunks += len(texts)
            logger.info("Indexed %s chunks from %s", len(texts), path.name)
        except Exception as exc:
            failed_files += 1
            logger.exception("Skipping unreadable/corrupted file %s: %s", path, exc)

    summary = {
        "indexed_files": indexed_files,
        "skipped_files": skipped_files,
        "failed_files": failed_files,
        "total_chunks": total_chunks,
        "collection": COLLECTION_NAME,
        "chroma_path": str(CHROMA_DIR),
    }
    logger.info("Ingestion complete: %s", summary)
    return summary


def main() -> None:
    parser = argparse.ArgumentParser(description="Index Notun Alo RAG documents into ChromaDB.")
    parser.add_argument("--rebuild", action="store_true", help="Clear and rebuild the ChromaDB collection.")
    parser.add_argument("--file", type=str, default=None, help="Index one uploaded/specific file only.")
    args = parser.parse_args()
    summary = ingest_documents(rebuild=args.rebuild, only_file=Path(args.file) if args.file else None)
    print(summary)


if __name__ == "__main__":
    main()
